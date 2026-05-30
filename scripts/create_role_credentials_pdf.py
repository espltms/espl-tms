from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "espl-tms-login-credentials.docx"

PASSWORD = "TMSAdminPassword2026!"

ROWS = [
    ("Super Admin", "superadmin@espl.com", PASSWORD, "Complete access to all regions and all modules."),
    ("Place-1 Admin", "place1.admin@espl.com", PASSWORD, "Region admin access for Place-1."),
    ("Place-2 Admin", "place2.admin@espl.com", PASSWORD, "Region admin access for Place-2."),
    ("Place-3 Admin", "place3.admin@espl.com", PASSWORD, "Region admin access for Place-3."),
    ("Place-4 Admin", "place4.admin@espl.com", PASSWORD, "Region admin access for Place-4."),
    ("Place-5 Admin", "place5.admin@espl.com", PASSWORD, "Region admin access for Place-5."),
    ("Vendor 1", "vendor1@espl.com", PASSWORD, "Vendor access: vehicle summary only for Vendor 1."),
    ("Vendor 2", "vendor2@espl.com", PASSWORD, "Vendor access: vehicle summary only for Vendor 2."),
    ("Vendor 3", "vendor3@espl.com", PASSWORD, "Vendor access: vehicle summary only for Vendor 3."),
]


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_cell_border(cell, color="D9E2EC"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_text(cell, text, bold=False, color="1E293B", size=9):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_border(cell)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.65)
section.right_margin = Inches(0.65)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(10)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("ESPL TMS Login Credentials")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(11, 37, 69)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run("New access roles: Super Admin, Place Admins, and Vendors")
subtitle_run.font.size = Pt(10)
subtitle_run.font.color.rgb = RGBColor(100, 116, 139)

note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.LEFT
note_run = note.add_run("Important: Self registration is disabled. Accounts must be created/managed by the Super Admin.")
note_run.bold = True
note_run.font.size = Pt(10)
note_run.font.color.rgb = RGBColor(154, 52, 18)

table = doc.add_table(rows=1, cols=4)
table.autofit = False
widths = [Inches(1.35), Inches(2.0), Inches(1.55), Inches(2.65)]
headers = ["Role / Account", "Login Email", "Password", "Access"]
for idx, header in enumerate(headers):
    cell = table.rows[0].cells[idx]
    cell.width = widths[idx]
    set_cell_fill(cell, "E8F1FF")
    set_cell_text(cell, header, bold=True, color="0F172A", size=9)

for role, email, password, access in ROWS:
    row = table.add_row()
    values = [role, email, password, access]
    for idx, value in enumerate(values):
        cell = row.cells[idx]
        cell.width = widths[idx]
        set_cell_fill(cell, "FFFFFF" if len(table.rows) % 2 else "F8FAFC")
        set_cell_text(cell, value, bold=idx == 0, size=8 if idx in (1, 2, 3) else 9)

doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer.add_run("Keep this document private. Do not share vendor/admin passwords publicly.")
footer_run.font.size = Pt(9)
footer_run.font.color.rgb = RGBColor(100, 116, 139)

doc.save(OUT)
print(OUT)
